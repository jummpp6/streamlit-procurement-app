import os
from docx import Document


def replace_text_in_runs(paragraph, old_text, new_text):
    """ฟังก์ชันแทนที่ข้อความโดยรักษารูปแบบฟอนต์และตัวหนา/ตัวเอียงเดิมไว้ 100%"""
    if old_text in paragraph.text:
        # แทนที่ในกรณีที่ข้อความอยู่ใน run เดียวกัน
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return

        # ในกรณีที่ข้อความถูกแยกเป็นหลาย runs ให้แทนที่ระดับย่อหน้า
        full_text = paragraph.text.replace(old_text, new_text)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            for run in paragraph.runs:
                run.text = ""
            first_run.text = full_text


def process_docx(file_path, replacements):
    doc = Document(file_path)

    # 1. แทนที่ในย่อหน้าปกติ
    for p in doc.paragraphs:
        for old_t, new_t in replacements.items():
            replace_text_in_runs(p, old_t, new_t)

    # 2. แทนที่ในตาราง
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_t, new_t in replacements.items():
                        replace_text_in_runs(p, old_t, new_t)

    # 3. แทนที่ใน ส่วนหัว/ส่วนท้าย กระดาษ (Header/Footer)
    for section in doc.sections:
        for header_p in section.header.paragraphs:
            for old_t, new_t in replacements.items():
                replace_text_in_runs(header_p, old_t, new_t)
        for footer_p in section.footer.paragraphs:
            for old_t, new_t in replacements.items():
                replace_text_in_runs(footer_p, old_t, new_t)

    return doc


# ==========================================
# 1. กรอกข้อมูลโครงการใหม่ตรงนี้ที่เดียว
# ==========================================
data_input = {
    "{{PROJECT_NAME}}": " ",
    "{{BUDGET}}": " ",
    "{{BUDGET_TEXT}}": " ",
    "{{VENDOR_NAME}}": " ",
    "{{DOC_DATE}}": " ",
    "{{DOC_NO}}": " ",
    "{{PROJECT_NAME}}": " ",
    "{{BUDGET}}": " ",
    "{{BUDGET_MID}}": " ",
    "{{BUDGET_TEXT}}": " ",
    "{{BUDGET_WITH_TEXT}}": " ",
    "{{BUDGET_WITH_TEXT_MID}}": " ",
    "{{BUDGET_TYPE}}": " ",
    "{{ITEM_COUNT}}": " ",
    "{{DEPARTMENT}}": " ",
    "{{VENDOR_NAME}}": " ",
    "{{DOC_DATE}}": " ",
    "{{DOC_DATE2}}": " ",
    "{{DOC_DATE3}}": " ",
    "{{DOC_DATE4}}": " ",
    "{{DOC_NO}}": " ",
    "{{DOC_NO_SAVE}}": " ",
    "{{DOC_NO_SAVE2}}": " ",
    "{{DOC_NO_SAVE3}}": " ",
    "{{DOC_NO_SUBMIT}}": " ",
    "{{DIRECTOR_NAME_BUY}}": " ",
    "{{DIRECTOR_NAME_BUY2}}": " ",
    "{{DIRECTOR_NAME_BUY3}}": " ",
    "{{CHECKITEM_NAME}}": " ",
    "{{CHECKITEM_NAME2}}": " ",
    "{{CHECKITEM_NAME3}}": " ",
}

# ==========================================
# 2. รันระบบสร้างเอกสารทั้งหมด
# ==========================================
output_dir = "ผลลัพธ์เอกสาร"
os.makedirs(output_dir, exist_ok=True)

# ดึงไฟล์ .docx ทั้งหมดในโฟลเดอร์ (ยกเว้นไฟล์ที่สร้างออกมาแล้ว)
docx_files = [
    f
    for f in os.listdir(".")
    if f.endswith(".docx") and not f.startswith("~$")
]

for file_name in docx_files:
    print(f"กำลังประมวลผล: {file_name} ...")
    processed_doc = process_docx(file_name, data_input)

    # บันทึกไฟล์ใหม่ไปยังโฟลเดอร์ผลลัพธ์
    save_path = os.path.join(output_dir, f"เสร็จ_{file_name}")
    processed_doc.save(save_path)

print("\nเสร็จสิ้น! เอกสารทั้งหมดถูกสร้างไว้ในโฟลเดอร์ 'ผลลัพธ์เอกสาร'")
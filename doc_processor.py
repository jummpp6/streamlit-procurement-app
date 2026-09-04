# ==========================================
# ไฟล์: doc_processor.py (ฉบับสมบูรณ์: ใช้ระบบ Block ครบวงจร)
# ==========================================
import io
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def to_thai_num(text):
    if not text:
        return ""
    arabic_digits = "0123456789"
    thai_digits = "๐๑๒๓๔๕๖๗๘๙"
    return str(text).translate(str.maketrans(arabic_digits, thai_digits))


def format_thai_date(date_obj, use_thai=True):
    if not date_obj:
        return ""

    THAI_MONTHS = [
        "",
        "มกราคม",
        "กุมภาพันธ์",
        "มีนาคม",
        "เมษายน",
        "พฤษภาคม",
        "มิถุนายน",
        "กรกฎาคม",
        "สิงหาคม",
        "กันยายน",
        "ตุลาคม",
        "พฤศจิกายน",
        "ธันวาคม",
    ]

    day = date_obj.day
    month = THAI_MONTHS[date_obj.month]
    year = date_obj.year + 543

    if use_thai:
        day_str = to_thai_num(day)
        year_str = to_thai_num(year)
    else:
        day_str = str(day)
        year_str = str(year)

    return f"{day_str}  {month}  พ.ศ. {year_str}"


def format_budget_money(text, use_thai=True):
    if not text:
        return ""
    clean_str = str(text).replace(",", "").replace(" ", "").replace("บาท", "").strip()
    try:
        val = float(clean_str)
        formatted = f"{int(val):,}" if val.is_integer() else f"{val:,.2f}"
    except ValueError:
        formatted = text
    return to_thai_num(formatted) if use_thai else formatted


def set_font_exact_16(run, font_name="TH SarabunPSK", font_size_pt=16, is_bold=None):
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rPr.append(rFonts)
    if is_bold is not None:
        run.bold = is_bold


def replace_text_in_paragraph(paragraph, replacements, default_font="TH SarabunPSK"):
    full_text = "".join([run.text for run in paragraph.runs])

    # ลบ Tag START/END ทุกรูปแบบที่อาจค้างอยู่
    full_text = re.sub(r"\{\{.*?(START|END)_.*?\}\}", "", full_text)

    has_target_key = any(key in full_text for key in replacements.keys())

    if not has_target_key:
        orig_text = "".join([run.text for run in paragraph.runs])
        if orig_text != full_text:
            first_run = paragraph.runs[0] if paragraph.runs else None
            is_bold = first_run.bold if first_run else False
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = full_text
                set_font_exact_16(
                    paragraph.runs[0], default_font, font_size_pt=16, is_bold=is_bold
                )
        return

    first_run = paragraph.runs[0] if paragraph.runs else None
    is_bold = first_run.bold if first_run else False

    for key, value in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, str(value))

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = full_text
        set_font_exact_16(
            paragraph.runs[0], default_font, font_size_pt=16, is_bold=is_bold
        )


def process_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_text_in_paragraph(p, replacements)
            for nested_table in cell.tables:
                process_table(nested_table, replacements)


def remove_row_from_table(table, keywords_to_remove):
    rows_to_delete = []
    for row in table.rows:
        row_text = "".join([cell.text for cell in row.cells])
        if any(keyword in row_text for keyword in keywords_to_remove):
            rows_to_delete.append(row)

    for row in rows_to_delete:
        tr = row._tr
        parent = tr.getparent()
        if parent is not None:
            parent.remove(tr)


def remove_block_by_tags(doc, start_tag, end_tag):
    """ฟังก์ชันลบบล็อกอัจฉริยะ: รองรับทั้งหน้ากระดาษและบล็อกบรรทัด ทั้งใน Body หลักและในตาราง"""

    def process_elements(parent_container):
        elements_to_remove = []
        inside_block = False

        for element in list(parent_container):
            text = "".join(element.itertext()) if hasattr(element, "itertext") else ""

            if start_tag in text:
                inside_block = True
                elements_to_remove.append(element)
                if end_tag in text:
                    inside_block = False
                continue

            if inside_block:
                elements_to_remove.append(element)
                if end_tag in text:
                    inside_block = False

        for element in elements_to_remove:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)

    # 1. ค้นหาและลบใน Body หลักของเอกสาร
    process_elements(doc.element.body)

    # 2. ค้นหาและลบในตารางทั้งหมด (รวมถึงตารางซ้อนเซลล์) เพื่อให้แน่ใจว่าบล็อกในตารางถูกลบเกลี้ยง
    def check_tables(tbl):
        for row in tbl.rows:
            for cell in row.cells:
                process_elements(cell._tc)
                for nested_tbl in cell.tables:
                    check_tables(nested_tbl)

    for table in doc.tables:
        check_tables(table)


def clean_unused_rows(doc, shop_count=1, buy_count=3, check_count=3):
    keywords_to_remove = []

    # 1. เช็กแท็กร้านค้าส่วนเกิน (ใช้ระบบ Block ทั้งแบบหน้ากระดาษ และแบบบรรทัดเฉพาะกิจ)
    for i in range(shop_count + 1, 5):
        keywords_to_remove.append(f"{{{{VENDOR_NAME{i}}}}}")
        keywords_to_remove.append(f"{{{{VENDOR_NAME_{i}}}}}")

        # ลบบล็อกหน้ากระดาษส่วนเกิน (เช่น ใบสั่งซื้อ, ใบตรวจรับ)
        remove_block_by_tags(doc, f"{{{{START_SHOP{i}}}}}", f"{{{{END_SHOP{i}}}}}")

        # ลบบล็อกบรรทัดเฉพาะกิจ (เช่น บันทึกสรุป ที่ครอบด้วย START_SHOP_LINE / END_SHOP_LINE)
        remove_block_by_tags(
            doc, f"{{{{START_SHOP_LINE{i}}}}}", f"{{{{END_SHOP_LINE{i}}}}}"
        )

    # 2. เช็กแท็กกรรมการจัดซื้อส่วนเกิน
    for i in range(buy_count + 1, 4):
        keywords_to_remove.append(f"{{{{DIRECTOR_NAME_BUY{i}}}}}")
        keywords_to_remove.append(f"{{{{SIGN_LINE_BUY{i}}}}}")

    # 3. เช็กแท็กกรรมการตรวจรับส่วนเกิน
    for i in range(check_count + 1, 4):
        keywords_to_remove.append(f"{{{{CHECKITEM_NAME{i}}}}}")
        keywords_to_remove.append(f"{{{{SIGN_LINE_CHECK{i}}}}}")

    if not keywords_to_remove:
        return

    for table in doc.tables:
        remove_row_from_table(table, keywords_to_remove)
        for row in table.rows:
            for cell in row.cells:
                for nested_table in cell.tables:
                    remove_row_from_table(nested_table, keywords_to_remove)


def remove_remaining_tags(doc):
    """ลบ Tag สัญลักษณ์ START/END ทุกรูปแบบที่อาจหลงเหลืออยู่ออกทั้งหมด"""
    pattern = re.compile(r"\{\{.*?(START|END)_.*?\}\}")

    def clean_runs(paragraphs):
        for p in paragraphs:
            if pattern.search(p.text):
                for run in p.runs:
                    if pattern.search(run.text):
                        run.text = pattern.sub("", run.text)

    clean_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                clean_runs(cell.paragraphs)


def remove_trailing_empty_paragraphs(doc):
    """ลบย่อหน้าว่างเปล่าท้ายไฟล์"""
    while doc.paragraphs:
        last_p = doc.paragraphs[-1]
        if not last_p.text.strip():
            p_xml = last_p._element
            p_xml.getparent().remove(p_xml)
        else:
            break


def process_docx(
    file_path, replacements_processed, shop_count=1, buy_count=3, check_count=3
):
    doc = Document(file_path)

    # 1. จัดการลบบล็อกส่วนเกินทั้งหมด (ทั้งหน้ากระดาษ และบรรทัดเฉพาะกิจ) ด้วยระบบ Block เดียวกัน
    clean_unused_rows(
        doc, shop_count=shop_count, buy_count=buy_count, check_count=check_count
    )

    # 2. แทนที่ข้อความในย่อหน้าปกติ
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, replacements_processed)

    # 3. แทนที่ข้อความในตาราง
    for table in doc.tables:
        process_table(table, replacements_processed)

    # 4. ทำความสะอาด Tag ที่อาจหลงเหลือ
    remove_remaining_tags(doc)

    # 5. คลีนย่อหน้าว่างท้ายไฟล์
    remove_trailing_empty_paragraphs(doc)

    # 6. ล็อกระยะขอบบน และเคลียร์ Header ไม่ให้ดันระยะขอบ
    for section in doc.sections:
        section.top_margin = Cm(1.25)
        section.header_distance = Cm(0)

        header = section.header
        for p in header.paragraphs:
            p.text = ""
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
